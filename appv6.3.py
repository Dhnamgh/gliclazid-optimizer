# Import thư viện
import streamlit as st
import pandas as pd
import numpy as np
import requests
import scipy.stats as stats
import shap
from io import BytesIO
import matplotlib.pyplot as plt
from uuid import uuid4
from sklearn.linear_model import LinearRegression, Lasso
from sklearn.ensemble import RandomForestRegressor
from sklearn.neural_network import MLPRegressor
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score, max_error
import statsmodels.formula.api as smf
from scipy.stats import shapiro
from scipy.optimize import differential_evolution
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import docx
from docx.shared import Inches
from PIL import Image
import tempfile
import io
import os
from docx import Document
from docx.shared import Inches
from PIL import Image

def format_number(val):
    try:
        if pd.isna(val):
            return ""
        if isinstance(val, (int, np.integer)):
            return str(val)
        if isinstance(val, (float, np.floating)):
            return f"{val:.0f}" if val.is_integer() else f"{val:.3f}"
        return str(val)
    except:
        return str(val)

from statsmodels.stats.diagnostic import het_breuschpagan, linear_reset
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.stattools import durbin_watson
import statsmodels.api as sm

# Hàm định dạng số: nếu là số nguyên thì không có chữ số thập phân, nếu là số thực thì có 3 chữ số
def auto_fmt(x):
    try:
        x_float = float(x)
        if x_float.is_integer():
            return f"{int(x_float)}"
        else:
            return f"{x_float:.3f}"
    except:
        return str(x)

# Chuẩn hoá dict best từ nhiều nguồn (optimal_row/best_formula)
def _as_best_dict(obj):
    """
    Chuẩn hóa về {'x1':..., 'x2':..., 'x3':..., 'y': (tùy)}
    Hỗ trợ:
      - dict có 'x1','x2','x3'
      - list/tuple/ndarray: [x1,x2,x3]
      - dict kiểu hiển thị: {"Primellose (%)":..., "PVP (%)":..., "Aerosil (%)":...}
    """
    if obj is None:
        return None
    # dạng đã chuẩn
    if isinstance(obj, dict) and all(k in obj for k in ("x1","x2","x3")):
        d = {'x1': float(obj['x1']), 'x2': float(obj['x2']), 'x3': float(obj['x3'])}
        if 'objective' in obj:
            d['y_pred'] = float(obj['objective'])
        return d
    # dạng hiển thị optimal_row
    if isinstance(obj, dict) and all(k in obj for k in ("Primellose (%)","PVP (%)","Aerosil (%)")):
        return {
            'x1': float(obj["Primellose (%)"]),
            'x2': float(obj["PVP (%)"]),
            'x3': float(obj["Aerosil (%)"]),
            'y_pred': float(obj.get("y (dự đoán)", np.nan)) if obj.get("y (dự đoán)") is not None else np.nan
        }
    # list/tuple/array
    if isinstance(obj, (list, tuple, np.ndarray)) and len(obj) >= 3:
        return {'x1': float(obj[0]), 'x2': float(obj[1]), 'x3': float(obj[2])}
    return None

# Hàm tạo phương trình hồi quy
def gen_regression_equation(model, response_name="y"):
    params = model.params
    terms = [f"{coef:.2f}*{var}" for var, coef in params.items() if var != "Intercept"]
    intercept = f"{params['Intercept']:.2f}"
    return f"{response_name} = {intercept} + " + " + ".join(terms)

# --- Auto-load default CSV once at startup ---
def _autoload_df_once():
    if st.session_state.get("df") is None:
        try:
            st.session_state["df"] = pd.read_csv("Gliclazid Data.csv")
        except FileNotFoundError:
            # Không chặn app; các tab sẽ cảnh báo nếu file không có
            st.session_state["df"] = None

_autoload_df_once()

# ===================== CẤU HÌNH STREAMLIT + LOGIN =====================
st.set_page_config(
    page_title="Gliclazid Optimizer",
    layout="wide",
    initial_sidebar_state="expanded"  # giúp sidebar dễ bật (nhất là trên mobile)
)
# Che badge/profile góc phải dưới + ẩn menu, nhưng CHO PHÉP hiện nút ≡/<< trên mobile
st.markdown("""
<style>
/* Ẩn menu hệ thống & footer để gọn gàng (desktop) */
#MainMenu { visibility: hidden; }
footer    { visibility: hidden; }

/* Mặc định ẩn header (trên desktop) */
header    { visibility: hidden; }

/* ✅ Trên mobile: BẬT LẠI header để có nút ≡/<< mở sidebar */
@media (max-width: 768px){
  header { visibility: visible !important; }
}

/* Ẩn mục nav multipage mặc định ở sidebar (tránh người dùng đi sang app khác) */
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] { display: none !important; }
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] a {
  pointer-events: none !important; opacity: .2 !important; cursor: not-allowed !important;
}

/* Lớp chắn click badge/profile ở góc phải dưới (áp dụng cho cả desktop & mobile) */
#st-bottom-blocker {
  position: fixed; right: 0; bottom: 0;
  width: 170px; height: 170px;
  z-index: 999999; pointer-events: auto;
}
#st-bottom-blocker .mask {
  position: absolute; inset: 0;
  background: transparent; /* đổi rgba(0,0,0,.05) nếu muốn mờ nhẹ */
}

/* Mobile: badge nhỏ hơn -> thu vùng chắn chút */
@media (max-width: 600px){
  #st-bottom-blocker { width: 130px; height: 130px; }
}
</style>

<!-- lớp chắn ở góc dưới phải -->
<div id="st-bottom-blocker"><div class="mask" title="blocked"></div></div>
""", unsafe_allow_html=True)

# --- Session defaults: luôn có dù bị reload trong iframe ---
def _ensure_defaults():
    if "df" not in st.session_state:
        st.session_state["df"] = None
    if "model" not in st.session_state:
        st.session_state["model"] = None
    # targets: suy luận từ df nếu có, nếu chưa thì tối thiểu có y1
    if "targets" not in st.session_state:
        df = st.session_state["df"]
        if isinstance(df, pd.DataFrame):
            ys = [c for c in df.columns if str(c).startswith("y")]
            st.session_state["targets"] = {y: y for y in ys} if ys else {"y1": "y1"}
        else:
            st.session_state["targets"] = {"y1": "y1"}

_ensure_defaults()

# CSS ẩn menu/header/footer và vô hiệu nav góc trái
_HARDEN_CSS = """
<style>
#MainMenu {visibility: hidden;}
header {visibility: hidden;}
footer {visibility: hidden;}
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] {display: none !important;}
section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] a {
  pointer-events: none !important; opacity: .2 !important; cursor: not-allowed !important;
}
* { -webkit-user-drag: none; }
body.app-locked [data-testid="stSidebar"] * { pointer-events:none; opacity:.35; }
</style>
"""
st.markdown(_HARDEN_CSS, unsafe_allow_html=True)

def _safe_rerun():
    """Tương thích nhiều phiên bản Streamlit."""
    try:
        st.rerun()  # Streamlit >= 1.29
    except AttributeError:
        st.experimental_rerun()  # Streamlit cũ

def _login_gate():
    """Form đăng nhập; trả True nếu đã đăng nhập."""
    if "auth_ok" not in st.session_state:
        st.session_state["auth_ok"] = False

    # Toggle lớp mờ khi chưa đăng nhập
    st.markdown(
        "<script>document.body.classList.%s('app-locked');</script>" %
        ("add" if not st.session_state["auth_ok"] else "remove"),
        unsafe_allow_html=True
    )

    if st.session_state["auth_ok"]:
        col1, col2 = st.columns([1,1])
        with col1:
            st.success(f"🔐 Đã đăng nhập: **{st.session_state.get('who','user')}**")
        with col2:
            if st.button("Đăng xuất"):
                st.session_state.clear()
                _safe_rerun()
        return True

    with st.container(border=True):
        st.subheader("🔐 Đăng nhập để sử dụng ứng dụng")
        with st.form("login_form", clear_on_submit=False):
            u = st.text_input("Tên đăng nhập", autocomplete="username")
            p = st.text_input("Mật khẩu", type="password", autocomplete="current-password")
            ok = st.form_submit_button("Đăng nhập")
        if ok:
            user_ok = u == st.secrets.get("APP_USERNAME", "")
            pass_ok = p == st.secrets.get("APP_PASSWORD", "")
            if user_ok and pass_ok:
                st.session_state["auth_ok"] = True
                st.session_state["who"] = u or "user"
                _safe_rerun()
            else:
                st.error("❌ Sai tên đăng nhập hoặc mật khẩu.")
    return False

# CSS giao diện và tiêu đề (an toàn, sau khi _HARDEN_CSS đã hợp lệ)
st.markdown("""
<style>
body, div, p { font-family: 'Open Sans', sans-serif; font-size:15px; color:#333; }
.stButton>button { background-color:#007ac1; color:white; border-radius:6px; padding:8px 16px; font-weight:bold; border:none; }
.stButton>button:hover { background-color:#045c87; }
</style>
<div style='background-color:#f4f8fb; padding:25px; border-radius:12px; text-align:center; margin-bottom:30px;'>
  <div style='font-size:32px; font-weight:bold; color:#045c87;'>Ứng dụng AI trong tối ưu hoá Gliclazid</div>
  <div style='font-size:16px; color:#666;'>Thiết kế công thức tá dược tối ưu</div>
</div>
""", unsafe_allow_html=True)

# ===================== HẾT KHỐI ĐẦU =====================


# 📌 Sidebar điều hướng
st.sidebar.image("background.png", use_container_width=True)
st.sidebar.title("Gliclazid Optimizer")
# 🔐 BẮT BUỘC ĐĂNG NHẬP TRƯỚC
if not _login_gate():
    st.stop()

tab = st.sidebar.radio("🔍 Chọn chức năng", [
    "📤 Dữ liệu", "🧩 Trực quan hóa dữ liệu", "🧮 Phân tích độ nhạy", "📊 Mô hình", "🧠 Diễn giải mô hình", "📈 Thống kê mô tả", "📉 Kiểm định", "🎯 Tối ưu", "📝 Báo cáo",
    "📄 Phân tích hồi quy", "🔗 So sánh các mô hình", "📤 Xuất kết quả", "📬 Phản hồi"
])

# Tab Dữ liệu
if tab == "📤 Dữ liệu":
    try:
        # Đọc sẵn file CSV có sẵn trong cùng thư mục với file app
        df = pd.read_csv("Gliclazid Data.csv")

        # Lưu vào session_state để các tab khác dùng
        st.session_state.df = df

        st.success("✅ Đã nạp dữ liệu mẫu: 'Gliclazid Data.csv'")
        st.dataframe(df.style.format(format_number))

        # (Tùy chọn) Cho phép người dùng tải file khác nếu muốn
        with st.expander("📁 Tải file CSV khác (tùy chọn)"):
            uploaded_file = st.file_uploader("Chọn file CSV khác:", type=["csv"])
            if uploaded_file is not None:
                df_new = pd.read_csv(uploaded_file)
                st.session_state.df = df_new
                st.info("📄 Dữ liệu mới đã được thay thế thành công.")
                st.dataframe(df_new.style.format(format_number))

    except FileNotFoundError:
        st.error("❌ Không tìm thấy file 'Gliclazid Data.csv' trong cùng thư mục ứng dụng.")
        st.stop()


# Trực quan hóa dữ liệu
if tab == "🧩 Trực quan hóa dữ liệu":
    st.header("🧩 So sánh biểu đồ 2 biến định lượng")

    df = st.session_state.get("df")
    if df is None:
        st.warning("📂 Chưa có dữ liệu.")
        st.stop()

    numeric_cols = df.select_dtypes(include=["float", "int"]).columns.tolist()

    col1, col2 = st.columns(2)
    with col1:
        var1 = st.selectbox("🔸 Chọn biến thứ nhất", numeric_cols, key="var1")
    with col2:
        var2 = st.selectbox("🔹 Chọn biến thứ hai", numeric_cols, key="var2")

    chart_type = st.radio("📈 Chọn loại biểu đồ", ["Histogram", "Boxplot", "Density"], horizontal=True)

    def render_chart(column_name, chart_type, ax):
        if chart_type == "Histogram":
            sns.histplot(df[column_name], kde=True, ax=ax, color='skyblue')
        elif chart_type == "Boxplot":
            sns.boxplot(y=df[column_name], ax=ax, color='lightgreen')
        elif chart_type == "Density":
            sns.kdeplot(df[column_name], ax=ax, fill=True, color='orange')
        ax.set_title(f"{chart_type} of {column_name}")
        ax.set_xlabel(column_name)
        ax.set_ylabel("Tần suất" if chart_type == "Histogram" else "")

    # Tạo biểu đồ kép
    fig_all, axes = plt.subplots(1, 2, figsize=(10, 4))
    render_chart(var1, chart_type, axes[0])
    render_chart(var2, chart_type, axes[1])
    fig_all.tight_layout()

    # Hiển thị lên Streamlit
    st.pyplot(fig_all)

    # Lưu vào session_state để dùng trong tab Xuất kết quả
    st.session_state["eda_plot"] = fig_all

    # Tải hình
    buf = BytesIO()
    fig_all.savefig(buf, format="png")
    buf.seek(0)
    st.download_button(
        label="📥 Tải biểu đồ so sánh",
        data=buf,
        file_name=f"Comparison_{var1}_vs_{var2}.png",
        mime="image/png"
    )

# Tab Phân tích độ nhạy
if tab == "🧮 Phân tích độ nhạy":
    st.header("🧮 Phân tích độ nhạy (Sensitivity Analysis)")
    
    df = st.session_state.get("df")
    if df is None:
        st.warning("📂 Vui lòng tải dữ liệu ở tab 📤 Dữ liệu.")
        st.stop()

    # Biến đầu vào và đầu ra
    input_cols = ["x1", "x2", "x3"]
    target_candidates = [col for col in df.columns if col.startswith("y")]
    
    if not all(col in df.columns for col in input_cols) or not target_candidates:
        st.warning("⚠️ Dữ liệu cần có các cột x1, x2, x3 và ít nhất một cột y.")
        st.stop()

    # Chọn biến mục tiêu và thêm key duy nhất
    target_col = st.selectbox("🎯 Chọn biến mục tiêu để phân tích", target_candidates, key="target_select_sensitivity")
    X = df[input_cols]
    y = df[target_col]

    # Huấn luyện mô hình Random Forest
    model = RandomForestRegressor(n_estimators=100, random_state=42)
    model.fit(X, y)
    importances = model.feature_importances_

    # Tạo bảng dữ liệu
    importance_df = pd.DataFrame({
        "Biến": input_cols,
        "Tầm quan trọng": importances
    }).sort_values("Tầm quan trọng", ascending=False)

    st.markdown(f"📌 **Tầm quan trọng của các biến đầu vào đối với `{target_col}`:**")
    st.dataframe(importance_df.style.format(format_number), use_container_width=True)

    # 🔁 Gán vào session_state để xuất ra Word
    st.session_state["importance_df"] = importance_df

    # Vẽ biểu đồ
    fig, ax = plt.subplots(figsize=(6, 4))
    sns.barplot(x="Tầm quan trọng", y="Biến", data=importance_df, ax=ax, palette="Blues_d")
    ax.set_title(f"Feature Importance - {target_col}", fontsize=12)
    ax.set_xlabel("Tầm quan trọng", fontsize=12)
    ax.set_ylabel("Biến đầu vào", fontsize=12)
    ax.tick_params(axis='both', labelsize=12)
    fig.tight_layout()
    st.pyplot(fig)

    # 🔁 Gán vào session_state để xuất ra Word
    st.session_state["importance_fig"] = fig

    # Tải biểu đồ PNG, dùng uuid4 để tránh lỗi trùng key
    buf = BytesIO()
    fig.savefig(buf, format="png")
    buf.seek(0)

    st.download_button(
        label="📥 Tải biểu đồ PNG",
        data=buf,
        file_name=f"feature_importance_{target_col}.png",
        mime="image/png",
        key=f"download_importance_{target_col}_{uuid4()}"
    )

# Tab 📊 Mô hình
if tab == "📊 Mô hình":
    st.header("📊 Huấn luyện mô hình hồi quy")

    df = st.session_state.get("df")
    if df is None or df.empty:
        st.warning("📂 Vui lòng tải dữ liệu trước.")
        st.stop()

    required_cols = ['x1', 'x2', 'x3', 'y1']
    if not all(col in df.columns for col in required_cols):
        st.warning("⚠️ Dữ liệu không đầy đủ: cần có các cột x1, x2, x3, y1.")
        st.stop()

    X = df[['x1', 'x2', 'x3']]
    y = df['y1']

    st.subheader("🔧 Chọn mô hình hồi quy")
    model_type = st.selectbox("Chọn mô hình:", ["Linear Regression", "Lasso", "Random Forest", "Neural Network"])

    if model_type == "Linear Regression":
        model = LinearRegression()
    elif model_type == "Lasso":
        alpha = st.slider("Chọn alpha (Lasso):", 0.01, 1.0, 0.1)
        model = Lasso(alpha=alpha)
    elif model_type == "Random Forest":
        n_trees = st.slider("Số cây trong rừng:", 10, 200, 100)
        model = RandomForestRegressor(n_estimators=n_trees, random_state=42)
    else:
        model = MLPRegressor(hidden_layer_sizes=(100,), max_iter=1000, random_state=42)

    model.fit(X, y)
    y_pred = model.predict(X)

    r2 = r2_score(y, y_pred)
    rmse = np.sqrt(mean_squared_error(y, y_pred))
    mae = mean_absolute_error(y, y_pred)

    st.success(f"✅ Huấn luyện hoàn tất với mô hình **{model_type}**")
    st.markdown(f"**R²:** {r2:.3f} &nbsp;&nbsp; | &nbsp;&nbsp; **RMSE:** {rmse:.3f} &nbsp;&nbsp; | &nbsp;&nbsp; **MAE:** {mae:.3f}")

    # 📌 Lưu mô hình để dùng ở các tab sau
    st.session_state["model"] = model

    # 📌 Tóm tắt mô hình
    model_summary = f"""Mô hình: {model_type}
- R²: {r2:.3f}
- RMSE: {rmse:.3f}
- MAE: {mae:.3f}"""
    st.text_area("📄 Tóm tắt mô hình huấn luyện:", model_summary, height=150)
    st.session_state["model_summary"] = model_summary

    # 📊 Biểu đồ phần dư
    residuals = y - y_pred
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(y_pred, residuals, alpha=0.7)
    ax.axhline(0, color='red', linestyle='--')
    ax.set_xlabel("Giá trị dự đoán")
    ax.set_ylabel("Phần dư")
    ax.set_title("Biểu đồ phần dư")
    st.pyplot(fig)

    buf = BytesIO()
    fig.savefig(buf, format="png")
    buf.seek(0)
    st.session_state["residual_plot"] = buf

# Tab 🧠 Diễn giải mô hình
if tab == "🧠 Diễn giải mô hình":
    st.header("🧠 Diễn giải mô hình (SHAP)")

    df = st.session_state.get("df")
    model = st.session_state.get("model")

    if df is None or model is None:
        st.warning("⚠️ Cần tải dữ liệu và huấn luyện mô hình ở tab 📊 Mô hình trước.")
        st.stop()

    # Xác định cột đầu vào (x1,x2,x3). Có thể mở rộng nếu bạn có nhiều x*
    input_cols = [c for c in df.columns if c.startswith("x")]
    if not input_cols:
        st.warning("⚠️ Không tìm thấy cột đầu vào (x1, x2, x3, ...).")
        st.stop()

    X = df[input_cols]

    # Tạo explainer phù hợp (Explainer tự chọn backend: Tree, Linear, Kernel...)
    try:
        explainer = shap.Explainer(model, X)
        shap_values = explainer(X)
    except Exception as e:
        st.error(f"Không tính được SHAP cho mô hình này: {e}")
        st.stop()

    st.subheader("🔎 Mức độ ảnh hưởng tổng quát (bar)")
    plt.figure()
    shap.summary_plot(shap_values, X, plot_type="bar", show=False)
    fig_bar = plt.gcf()
    st.pyplot(fig_bar)

    st.subheader("🐝 Beeswarm (phân bố ảnh hưởng theo từng mẫu)")
    plt.figure()
    shap.summary_plot(shap_values, X, show=False)
    fig_bee = plt.gcf()
    st.pyplot(fig_bee)

    # Lưu 1 hình vào session_state để xuất báo cáo Word
    st.session_state["shap_plot"] = fig_bar

# Thống kê mô tả
if tab == "📈 Thống kê mô tả":
    df = st.session_state.get("df")
    if df is None:
        st.warning("⚠️ Bạn cần tải dữ liệu ở tab 📤 Dữ liệu.")
        st.stop()

    st.subheader("📈 Thống kê mô tả dữ liệu")

    # Hàm thống kê từng biến
    def descriptive_stats(series):
        series_clean = series.dropna()
        n = len(series_clean)
        mean = series_clean.mean()
        std = series_clean.std()
        se = std / np.sqrt(n) if n > 0 else np.nan
        ci_low, ci_high = stats.norm.interval(0.95, loc=mean, scale=se) if n > 1 else (np.nan, np.nan)

        shapiro_p = stats.shapiro(series_clean)[1] if 3 <= n <= 5000 else np.nan
        ks_p = stats.kstest(series_clean, 'norm', args=(mean, std))[1] if n > 0 else np.nan

        return {
            'N': n,
            'Missing': int(series.isna().sum()),
            'Mean': mean,
            'Std. Error of Mean': se,
            'CI 95% Lower': ci_low,
            'CI 95% Upper': ci_high,
            'Median': series_clean.median(),
            'Mode': series_clean.mode().iloc[0] if not series_clean.mode().empty else np.nan,
            'Std. Deviation': std,
            'Variance': series_clean.var(),
            'Skewness': series_clean.skew(),
            'Kurtosis': series_clean.kurt(),
            'Range': series_clean.max() - series_clean.min(),
            'Minimum': series_clean.min(),
            'Maximum': series_clean.max(),
            '25%': series_clean.quantile(0.25),
            '50%': series_clean.quantile(0.50),
            '75%': series_clean.quantile(0.75),
            'Shapiro-Wilk (p)': shapiro_p,
            'Kolmogorov-Smirnov (p)': ks_p
        }

    # Lấy các cột định lượng
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    result = {}

    for col in numeric_cols:
        result[col] = descriptive_stats(df[col])

    # Chuyển về dạng đúng: chỉ số là hàng, biến là cột
    stats_df = pd.DataFrame(result)

    # Hàm định dạng thông minh
    def auto_fmt2(val):
        try:
            val = float(val)
            return f"{val:.0f}" if val.is_integer() else f"{val:.3f}"
        except:
            return str(val)

    # Tạo bản đã định dạng để hiển thị đẹp
    formatted_df = stats_df.copy()
    for col in formatted_df.columns:
        formatted_df[col] = formatted_df[col].apply(auto_fmt2)

    # Hiển thị bảng đã định dạng
    st.dataframe(formatted_df, use_container_width=True)

    # Ghi vào session_state để xuất báo cáo
    st.session_state["stats_df"] = stats_df            # bản gốc
    st.session_state["stats_df_fmt"] = formatted_df    # bản định dạng

#Tab Kiểm định
if tab == "📉 Kiểm định":
    import statsmodels.api as sm
    from statsmodels.stats.diagnostic import het_breuschpagan, linear_reset
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    from statsmodels.stats.stattools import durbin_watson

    df = st.session_state.get("df")
    model = st.session_state.get("model")
    targets = (st.session_state.get("targets")
               or ({c: c for c in df.columns if str(c).startswith("y")} if isinstance(df, pd.DataFrame) else {"y1":"y1"}))

    if df is None or model is None:
        st.warning("⚠️ Cần tải dữ liệu và huấn luyện mô hình ở tab 📊 Mô hình trước.")
        st.stop()


    st.markdown("## 🧪 Kiểm định giả định hồi quy")

    X = df[["x1", "x2", "x3"]]
    X_const = sm.add_constant(X)

    regression_tests = []

    for target, label in targets.items():
        st.subheader(f"🎯 Biến đầu ra: {label} ({target})")

        y = df[target]
        model.fit(X, y)
        y_pred = model.predict(X)
        residuals = y - y_pred

        # Kiểm định phân phối chuẩn (Shapiro-Wilk)
        sw_stat, sw_p = shapiro(residuals)
        sw_ketluan = (
            "✅ Phần dư có phân phối chuẩn (p > 0.05)"
            if sw_p > 0.05 else "⚠️ Phần dư không tuân theo phân phối chuẩn (p ≤ 0.05)"
        )
        st.write(f"**Shapiro-Wilk:** W = `{sw_stat:.3f}`, p = `{sw_p:.4f}`")
        st.info(sw_ketluan)

        # Kiểm định phương sai không đổi (Breusch–Pagan)
        ols_model = sm.OLS(y, X_const).fit()
        bp_test = het_breuschpagan(ols_model.resid, X_const)
        bp_stat, bp_p = bp_test[0], bp_test[1]
        bp_ketluan = (
            "✅ Không có bằng chứng về phương sai thay đổi (p > 0.05)"
            if bp_p > 0.05 else "⚠️ Có thể có hiện tượng phương sai thay đổi (p ≤ 0.05)"
        )
        st.write(f"**Breusch–Pagan:** LM = `{bp_stat:.3f}`, p = `{bp_p:.4f}`")
        st.info(bp_ketluan)

        # Kiểm định tự tương quan phần dư (Durbin-Watson)
        dw_stat = durbin_watson(ols_model.resid)
        if dw_stat < 1.5:
            dw_ketluan = "⚠️ Có thể có tự tương quan dương."
        elif dw_stat > 2.5:
            dw_ketluan = "⚠️ Có thể có tự tương quan âm."
        else:
            dw_ketluan = "✅ Không có bằng chứng rõ ràng về tự tương quan."
        st.write(f"**Durbin–Watson:** DW = `{dw_stat:.3f}`")
        st.info(dw_ketluan)

        # Kiểm định Linearity (Ramsey RESET)
        reset_test = linear_reset(ols_model, power=2, use_f=True)
        reset_stat, reset_p = reset_test.fvalue, reset_test.pvalue
        reset_ketluan = (
            "✅ Mô hình tuyến tính là phù hợp (p > 0.05)"
            if reset_p > 0.05 else "⚠️ Có dấu hiệu mô hình không tuyến tính (p ≤ 0.05)"
        )
        st.write(f"**Ramsey RESET:** F = `{reset_stat:.3f}`, p = `{reset_p:.4f}`")
        st.info(reset_ketluan)

        # Tính hệ số VIF để kiểm tra đa cộng tuyến
        vif_data = pd.DataFrame()
        vif_data["Biến"] = X.columns
        vif_data["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
        vif_data["Kết luận"] = vif_data["VIF"].apply(
            lambda x: "⚠️ Cao (đa cộng tuyến)" if x > 10 else ("⚠️ Trung bình" if x > 5 else "✅ Chấp nhận được")
        )
        st.write("**🎯 VIF (Đa cộng tuyến):**")
        st.dataframe(vif_data.set_index("Biến"))

        # Lưu kết quả
        regression_tests.append({
            "Biến đầu ra": f"{label} ({target})",
            "Shapiro-Wilk": f"W = {sw_stat:.3f}, p = {sw_p:.4f} | {sw_ketluan}",
            "Breusch–Pagan": f"LM = {bp_stat:.3f}, p = {bp_p:.4f} | {bp_ketluan}",
            "Durbin-Watson": f"DW = {dw_stat:.3f} | {dw_ketluan}",
            "Ramsey RESET": f"F = {reset_stat:.3f}, p = {reset_p:.4f} | {reset_ketluan}",
            "VIF": vif_data.to_dict(orient="records")
        })

    # Lưu vào session_state
    st.session_state["regression_tests"] = regression_tests

# Tab Tối ưu
if tab == "🎯 Tối ưu":
    st.header("📌 Tối ưu công thức đầu vào để tối đa hóa đầu ra")

    df = st.session_state.get("df")
    model = st.session_state.get("model")

    if df is None or model is None:
        st.warning("⚠️ Cần có dữ liệu và mô hình huấn luyện trước.")
        st.stop()

    # Giới hạn vùng giá trị của từng biến đầu vào
    x1_min, x1_max = float(df["x1"].min()), float(df["x1"].max())
    x2_min, x2_max = float(df["x2"].min()), float(df["x2"].max())
    x3_min, x3_max = float(df["x3"].min()), float(df["x3"].max())

    def objective(x):
        x_input = np.array(x).reshape(1, -1)
        return -model.predict(x_input)[0]  # tối đa hóa y => tối thiểu hóa -y

    bounds = [(x1_min, x1_max), (x2_min, x2_max), (x3_min, x3_max)]
    result = differential_evolution(objective, bounds)

    best = result.x
    best_output = -result.fun

    st.success("✅ Đã tìm được công thức tối ưu đầu vào:")
    st.markdown(f"""
    - Primellose: `{best[0]:.2f}` %
    - PVP: `{best[1]:.2f}` %
    - Aerosil: `{best[2]:.2f}` %
    - 🔼 Dự đoán đầu ra tối ưu (y1): `{best_output:.2f}`
    """)

    # 🔁 Lưu vào session_state theo 2 dạng (để tương thích các tab khác)
    st.session_state["optimal_row"] = {
        "Primellose (%)": round(best[0], 2),
        "PVP (%)": round(best[1], 2),
        "Aerosil (%)": round(best[2], 2),
        "y (dự đoán)": round(best_output, 2)
    }
    st.session_state["best_formula"] = {
        "x1": float(best[0]),
        "x2": float(best[1]),
        "x3": float(best[2]),
        "objective": float(best_output)
    }

# Tab Báo cáo
if tab == "📝 Báo cáo":
    df = st.session_state.get("df")
    results = st.session_state.get("results", {})
    targets = st.session_state.targets
    model = st.session_state.get("model")

    # Lấy best từ mọi nguồn
    best = _as_best_dict(st.session_state.get("best_formula"))
    if best is None:
        best = _as_best_dict(st.session_state.get("optimal_row"))

    if df is None or best is None or model is None:
        st.warning("⚠️ Bạn cần chạy mô hình và tối ưu công thức trước.")
        st.stop()

    X = df[["x1", "x2", "x3"]]
    # Nếu chưa có y_pred trong best, dự đoán bằng model hiện tại cho y1
    if "y_pred" not in best or pd.isna(best.get("y_pred")):
        best["y_pred"] = float(model.predict(np.array([[best["x1"], best["x2"], best["x3"]]]))[0])

    # --- Tạo đoạn mô tả tổng quan ---
    report_text = f"""--- BÁO CÁO PHÂN TÍCH ---
Tác giả: Đào Hồng Nam
Ngày phân tích: {datetime.today().strftime('%d-%m-%Y')}

🔬 Công thức tối ưu (đầu vào):
- Primellose: {auto_fmt(best['x1'])} %
- PVP: {auto_fmt(best['x2'])} %
- Aerosil: {auto_fmt(best['x3'])} %
→ Dự đoán đầu ra tối ưu (y1): {auto_fmt(best['y_pred'])}
"""

    # --- Kết quả từ tab Mô hình (nếu có) ---
    if results:
        report_text += "\n📊 Kết quả mô hình đang chọn:\n"
        for target, metrics in results.items():
            label = targets.get(target, target)
            r2 = metrics['r2']
            mae = metrics['mae']
            report_text += f"- {label} ({target}): R² = {auto_fmt(r2)}, MAE = {auto_fmt(mae)}\n"

    # --- Kết quả tất cả mô hình (huấn luyện nhanh trên toàn bộ dữ liệu) ---
    all_models = {
        "Linear Regression": LinearRegression(),
        "Lasso Regression": Lasso(alpha=0.1),
        "Random Forest": RandomForestRegressor(n_estimators=100, random_state=42),
        "ANN (Neural Network)": MLPRegressor(hidden_layer_sizes=(100,), max_iter=1000, random_state=42)
    }
    for model_name, mdl in all_models.items():
        report_text += f"\n📊 Kết quả mô hình: {model_name}\n"
        for target, label in targets.items():
            if target not in df.columns:
                continue
            y = df[target]
            mdl.fit(X, y)
            y_pred = mdl.predict(X)
            r2 = r2_score(y, y_pred)
            mae = mean_absolute_error(y, y_pred)
            report_text += f"- {label} ({target}): R² = {auto_fmt(r2)}, MAE = {auto_fmt(mae)}\n"

    # --- Hiển thị văn bản ---
    st.code(report_text, language="markdown")
    st.session_state["best_formula_text"] = report_text

    # --- Ghi ra file Word ---
    if st.button("📥 Tải báo cáo Word"):
        doc = Document()
        doc.add_heading("BÁO CÁO PHÂN TÍCH TỔNG HỢP", 0)
        doc.add_paragraph(f"Tác giả: Đào Hồng Nam\nNgày: {datetime.today().strftime('%d-%m-%Y')}")
        doc.add_heading("🔬 Công thức tối ưu", level=1)
        doc.add_paragraph(st.session_state.get("best_formula_text", "Chưa có dữ liệu."))

        # Thêm kết quả mô hình đang chọn (nếu có)
        if results:
            doc.add_heading("📊 Kết quả mô hình đang chọn", level=1)
            for target, metrics in results.items():
                label = targets.get(target, target)
                doc.add_paragraph(
                    f"{label} ({target}): R² = {auto_fmt(metrics['r2'])}, MAE = {auto_fmt(metrics['mae'])}"
                )

        # Lưu ra file buffer
        buf = io.BytesIO()
        doc.save(buf)
        st.download_button("📤 Tải báo cáo Word", buf.getvalue(), file_name="bao_cao_phan_tich.docx")

# Tab Phân tích hồi quy (đồng bộ đúng tên tab)
if tab == "📄 Phân tích hồi quy":
    st.header("📉 Phân tích hồi quy tuyến tính")

    df = st.session_state.get("df")
    if df is None:
        st.warning("⚠️ Cần tải dữ liệu trước.")
        st.stop()

    target_col = st.selectbox("🎯 Chọn biến đầu ra (y):", options=[col for col in df.columns if col.startswith("y")])
    input_cols = [col for col in df.columns if col.startswith("x")]

    X = df[input_cols]
    y = df[target_col]

    X_const = sm.add_constant(X)
    model_ols = sm.OLS(y, X_const).fit()

    summary_str = model_ols.summary().as_text()

    st.text("📄 Kết quả hồi quy:")
    st.text(summary_str)

    # 🔁 Lưu vào session_state để xuất Word
    st.session_state["regression_summary"] = summary_str

# Tab So sánh các mô hình
if tab == "🔗 So sánh các mô hình":
    st.markdown("## 🔗 So sánh các mô hình hồi quy")
    df = st.session_state.get("df")
    if df is None or df.empty:
        st.warning("📂 Không tìm thấy dữ liệu đã tải.")
        st.stop()

    required_cols = ['x1', 'x2', 'x3', 'y1']
    if not all(col in df.columns for col in required_cols):
        st.warning("⚠️ Dữ liệu không có đủ cột x1, x2, x3, y1.")
        st.stop()

    X = df[['x1', 'x2', 'x3']]
    y = df['y1']
    n, k = X.shape

    def auto_fmt3(val):
        return f"{val:.0f}" if float(val).is_integer() else f"{val:.3f}"

    models = {
        "Linear Regression": LinearRegression(),
        "Lasso Regression": Lasso(alpha=0.1),
        "Random Forest": RandomForestRegressor(n_estimators=100, random_state=42),
        "ANN (Neural Network)": MLPRegressor(hidden_layer_sizes=(100,), max_iter=1000, random_state=42)
    }

    results_tbl = []
    for name, mdl in models.items():
        mdl.fit(X, y)
        pred = mdl.predict(X)

        mse = mean_squared_error(y, pred)
        rmse = np.sqrt(mse)
        mae = mean_absolute_error(y, pred)
        mape = np.mean(np.abs((y - pred) / y)) * 100 if np.all(y != 0) else np.nan
        r2 = r2_score(y, pred)
        adj_r2 = 1 - (1 - r2) * (n - 1) / (n - k - 1)
        maxerr = max_error(y, pred)

        results_tbl.append({
            "Mô hình": name,
            "R²": auto_fmt3(r2),
            "Adj. R²": auto_fmt3(adj_r2),
            "MSE": auto_fmt3(mse),
            "RMSE": auto_fmt3(rmse),
            "MAE": auto_fmt3(mae),
            "MAPE (%)": auto_fmt3(mape) if not np.isnan(mape) else "NA",
            "Max Error": auto_fmt3(maxerr)
        })

    df_result = pd.DataFrame(results_tbl)
    st.dataframe(df_result, use_container_width=True)

    # 🔄 Lưu vào session_state để dùng cho báo cáo/xuất file
    st.session_state["model_comparison"] = df_result

# Tab Xuất kết quả
if tab == "📤 Xuất kết quả":
    st.header("📤 Xuất kết quả tổng hợp")
    doc = Document()
    doc.add_heading('BÁO CÁO PHÂN TÍCH TỔNG HỢP', 0)

    # 1. Thống kê mô tả
    doc.add_heading("1. Thống kê mô tả", level=1)
    if "stats_df" in st.session_state:
        df_stats = st.session_state["stats_df"]
        table = doc.add_table(rows=1, cols=len(df_stats.columns) + 1)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Tham số"
        for i, col in enumerate(df_stats.columns):
            hdr[i+1].text = col
        for idx, row in df_stats.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            for i, val in enumerate(row):
                try:
                    row_cells[i+1].text = f"{float(val):.3f}"
                except:
                    row_cells[i+1].text = str(val)
    else:
        doc.add_paragraph("Chưa có dữ liệu thống kê mô tả.")

    # 2. Trực quan hóa
    doc.add_heading("2. Trực quan hóa dữ liệu", level=1)
    if "eda_plot" in st.session_state:
        img_path = "eda_plot.png"
        st.session_state["eda_plot"].savefig(img_path, bbox_inches='tight')
        doc.add_picture(img_path, width=Inches(5))
        os.remove(img_path)
    else:
        doc.add_paragraph("Chưa có biểu đồ trực quan hóa.")

    # 3. Phân tích độ nhạy
    doc.add_heading("3. Phân tích độ nhạy", level=1)
    if "importance_df" in st.session_state:
        df_imp = st.session_state["importance_df"]
        table = doc.add_table(rows=1, cols=len(df_imp.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df_imp.columns):
            table.cell(0, i).text = str(col)
        for _, row in df_imp.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
    else:
        doc.add_paragraph("Chưa có dữ liệu độ nhạy.")

    # 4. Mô hình đã huấn luyện
    doc.add_heading("4. Mô hình đã huấn luyện", level=1)
    if "model_summary" in st.session_state:
        doc.add_paragraph(st.session_state["model_summary"])
    else:
        doc.add_paragraph("Chưa có mô hình được huấn luyện.")

    # 5. Diễn giải mô hình
    doc.add_heading("5. Diễn giải mô hình", level=1)
    if "shap_plot" in st.session_state:
        img_path = "shap_plot.png"
        st.session_state["shap_plot"].savefig(img_path, bbox_inches='tight')
        doc.add_picture(img_path, width=Inches(5))
        os.remove(img_path)
    else:
        doc.add_paragraph("Chưa có biểu đồ diễn giải mô hình.")

    # 6. Kiểm định giả định hồi quy
    doc.add_heading("6. Kiểm định giả định hồi quy", level=1)
    if "regression_tests" in st.session_state:
        for res in st.session_state["regression_tests"]:
            para = doc.add_paragraph()
            para.add_run(f"{res['Biến đầu ra']}:\n").bold = True
            para.add_run(f"  - {res['Shapiro-Wilk']}\n")
            para.add_run(f"  - {res['Breusch–Pagan']}\n")
            para.add_run(f"  - {res['Durbin-Watson']}\n")
            para.add_run(f"  - {res['Ramsey RESET']}\n")
    else:
        doc.add_paragraph("Chưa có kết quả kiểm định.")

    # 7. Tối ưu công thức
    doc.add_heading("7. Tối ưu công thức", level=1)
    best = _as_best_dict(st.session_state.get("best_formula")) or _as_best_dict(st.session_state.get("optimal_row"))
    if best is not None:
        doc.add_paragraph(
            f"Primellose: {auto_fmt(best['x1'])}% | PVP: {auto_fmt(best['x2'])}% | Aerosil: {auto_fmt(best['x3'])}%"
        )
        if "y_pred" in best and not pd.isna(best["y_pred"]):
            doc.add_paragraph(f"Dự đoán y1 tối ưu: {auto_fmt(best['y_pred'])}")
    else:
        doc.add_paragraph("Chưa có công thức tối ưu.")

    # 8. Phân tích hồi quy
    doc.add_heading("8. Phân tích hồi quy", level=1)
    if "regression_summary" in st.session_state:
        doc.add_paragraph(st.session_state["regression_summary"])
    else:
        doc.add_paragraph("Chưa có kết quả phân tích hồi quy.")

    # 9. So sánh các mô hình
    doc.add_heading("9. So sánh các mô hình", level=1)
    if "model_comparison" in st.session_state:
        df_cmp = st.session_state["model_comparison"]
        table = doc.add_table(rows=1, cols=len(df_cmp.columns))
        table.style = "Table Grid"
        for i, col in enumerate(df_cmp.columns):
            table.cell(0, i).text = str(col)
        for _, row in df_cmp.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
    else:
        doc.add_paragraph("Chưa có bảng so sánh mô hình.")

    # Xuất tập tin
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Tải xuống báo cáo Word",
        data=buffer,
        file_name="bao_cao_phan_tich.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Tab Phản hồi
if tab == "📬 Phản hồi":
    st.markdown("## 📬 Góp ý & Phản hồi")
    with st.form("feedback_form"):
        name = st.text_input("👤 Họ tên")
        email = st.text_input("📧 Email")
        feedback_type = st.selectbox("📌 Loại phản hồi", ["Góp ý", "Báo lỗi", "Yêu cầu mới"])
        feedback = st.text_area("✍️ Nội dung phản hồi")

        submitted = st.form_submit_button("Gửi phản hồi")
        if submitted:
            st.success("✅ Cảm ơn bạn đã phản hồi!")

            # Gửi qua API giả lập (bọc try/except để không làm dừng app)
            try:
                requests.post(st.secrets.get("EMAIL_API_URL",""), json={
                    "to": st.secrets.get("EMAIL_TO",""),
                    "subject": f"Phản hồi từ {name} ({feedback_type})",
                    "body": f"Email: {email}\nLoại: {feedback_type}\nNội dung:\n{feedback}"
                })
            except Exception as e:
                st.info("Thông tin đã được ghi nhận (không gửi API).")

            try:
                requests.post(st.secrets.get("SHEET_API_URL",""), json={
                    "name": name,
                    "email": email,
                    "type": feedback_type,
                    "content": feedback,
                    "timestamp": datetime.now().strftime("%d-%m-%Y %H:%M:%S")
                })
            except Exception:
                pass

# ---------- FOOTER  ----------

st.markdown("---")
st.markdown("© Bản quyền thuộc về TS. Đào Hồng Nam - Đại học Y Dược Thành phố Hồ Chí Minh.")







